"""
论文工具 — MD → .tex / .pdf / .docx

使用方式：
  1. 在 paper/ 下写 draft.md（参照 template/template.md）
  2. 完成内容后运行 md_to_pdf() 或 md_to_tex()
  3. Word 接口作为应急备用
"""
import os
import subprocess


def md_to_pdf(md_path, output_dir=None):
    """
    MD → PDF（依赖 Pandoc + LaTeX）
    pandoc draft.md -o draft.pdf --pdf-engine=xelatex
    """
    if output_dir is None:
        output_dir = os.path.dirname(md_path)
    pdf_path = os.path.join(output_dir,
                            os.path.splitext(os.path.basename(md_path))[0] + '.pdf')

    try:
        subprocess.run([
            'pandoc', md_path, '-o', pdf_path,
            '--pdf-engine=xelatex',
            '-V', 'mainfont=SimSun',
            '-V', 'CJKmainfont=SimSun',
        ], check=True)
        print(f'PDF saved: {pdf_path}')
    except FileNotFoundError:
        print('Pandoc not found. Install: https://pandoc.org')
    except subprocess.CalledProcessError as e:
        print(f'Pandoc failed: {e}')
    return pdf_path


def md_to_tex(md_path, output_dir=None):
    """MD → .tex（不编译，用于手动调整）"""
    if output_dir is None:
        output_dir = os.path.dirname(md_path)
    tex_path = os.path.join(output_dir,
                            os.path.splitext(os.path.basename(md_path))[0] + '.tex')

    try:
        subprocess.run([
            'pandoc', md_path, '-o', tex_path,
        ], check=True)
        print(f'TEX saved: {tex_path}')
    except FileNotFoundError:
        print('Pandoc not found.')
    except subprocess.CalledProcessError as e:
        print(f'Pandoc failed: {e}')
    return tex_path


def md_to_docx(md_path, output_dir=None):
    """MD → .docx（Word 应急出口）"""
    if output_dir is None:
        output_dir = os.path.dirname(md_path)
    docx_path = os.path.join(output_dir,
                             os.path.splitext(os.path.basename(md_path))[0] + '.docx')

    try:
        subprocess.run([
            'pandoc', md_path, '-o', docx_path,
        ], check=True)
        print(f'DOCX saved: {docx_path}')
    except FileNotFoundError:
        print('Pandoc not found.')
    except subprocess.CalledProcessError as e:
        print(f'Pandoc failed: {e}')
    return docx_path


def write_docx_fallback(title, sections, filename='paper.docx'):
    """
    纯 Python 应急生成 Word（不依赖 Pandoc）
    sections: [('## 标题', '正文内容'), ...]
    """
    from docx import Document
    from docx.shared import Pt

    doc = Document()
    p = doc.add_paragraph()
    run = p.add_run(title)
    run.font.size = Pt(22)
    run.bold = True

    for heading, text in sections:
        if heading.startswith('## '):
            level = 2
        elif heading.startswith('### '):
            level = 3
        else:
            level = 1
        h = heading.lstrip('# ')
        doc.add_heading(h, level=level)
        doc.add_paragraph(text)

    doc.save(filename)
    print(f'DOCX saved: {filename}')
    return filename
